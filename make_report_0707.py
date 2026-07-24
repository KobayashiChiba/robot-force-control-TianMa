"""生成 2026-07-07 工作日报 docx"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── 标题 ──
title = doc.add_heading('2026-07-07 工作日报', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

# ── 概述 ──
doc.add_heading('一、今日工作：圆柱拟合', level=1)
doc.add_paragraph(
    '基于81个实测接触点，用最小二乘法拟合两个正交圆柱（轴∥Y 和轴∥Z）的几何参数，'
    '评估实际加工偏差。Y圆柱截面在XZ平面，Z圆柱截面在XY平面，分别独立做最小二乘圆拟合，共4个参数。'
)

# ── 拟合结果表格 ──
doc.add_heading('二、拟合结果', level=1)
table1 = doc.add_table(rows=6, cols=3, style='Light Shading Accent 1')
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['参数', 'Y方向圆柱 (∥Y)', 'Z方向圆柱 (∥Z)']
data = [
    ['轴线 X₀', '51.497 mm', '72.499 mm'],
    ['轴线 Y₀', '—', '65.151 mm'],
    ['轴线 Z₀', '-39.700 mm', '—'],
    ['半径', '9.002 mm', '17.998 mm'],
    ['RMS残差', '0.003 mm', '0.085 mm'],
]
for j, h in enumerate(headers):
    cell = table1.cell(0, j)
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table1.cell(i+1, j).text = val

doc.add_paragraph('')
p = doc.add_paragraph('注：Y方向圆柱轴∥Y，无固定Y坐标。')
p.runs[0].italic = True

# ── 偏差分析表格 ──
doc.add_heading('三、偏差分析', level=1)
table2 = doc.add_table(rows=6, cols=4, style='Light Shading Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
h2 = ['参数', '设计值', '实测值', '偏差']
d2 = [
    ['Y圆柱半径', '10 mm', '9.002 mm', '-0.998 mm'],
    ['Z圆柱半径', '20 mm', '17.998 mm', '-2.002 mm'],
    ['Y圆柱轴线 X₀', '0 mm', '51.497 mm', '—'],
    ['Y圆柱轴线 Z₀', '0 mm', '-39.700 mm', '—'],
    ['Z圆柱轴线 X₀', '27 mm', '72.499 mm', '—'],
]
for j, h in enumerate(h2):
    cell = table2.cell(0, j)
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
for i, row in enumerate(d2):
    for j, val in enumerate(row):
        table2.cell(i+1, j).text = val

doc.add_paragraph('')
doc.add_paragraph(
    '拟合精度很高：Y圆柱RMS残差仅0.003mm，Z圆柱RMS残差0.085mm。'
    '与实际设计值对比，两个圆柱半径均偏小（Y圆柱约-1mm，Z圆柱约-2mm），'
    '且工件整体坐标系与设计坐标系有较大平移。'
)

# ── 图像 ──
doc.add_heading('四、可视化结果', level=1)
img_path = r'C:\Users\KCserver\projects\formal\机器人末端力控\code\fig_collage_2x2.png'
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(5.5))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph('图1：圆柱拟合可视化（2×2布局）')
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in caption.runs:
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(100, 100, 100)

# ── 输出文件 ──
doc.add_heading('五、输出文件', level=1)
files_list = [
    'code/fit_cylinders.py — 圆柱拟合脚本',
    'code/fig_collage_2x2.png — 4图拼合（用于日报）',
    'code/fig1_points_and_cylinders.png — 测量点+两圆柱3D',
    'code/fig2_measured_vs_curve.png — 测量点vs交线对比',
    'code/fig3_xz_projection.png — Y投影（XZ平面）',
    'code/fig4_xy_projection.png — Z投影（XY平面）',
]
for f in files_list:
    doc.add_paragraph(f, style='List Bullet')

# ── 保存 ──
out = r'C:\Users\KCserver\projects\formal\机器人末端力控\2026-07-07_工作日报（圆柱拟合）.docx'
doc.save(out)
print(f'Saved: {out}')
