from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

doc.add_heading('机器人末端力控项目 — 工作日报', level=0)
doc.add_paragraph('日期：2026-07-03    状态：项目启动')

doc.add_heading('项目摘要', level=1)
doc.add_paragraph(
    '针对工业机器人末端力控应用中的圆柱孔相交边缘打磨路径规划问题，'
    '建立两正交圆柱交线的通用几何模型，并计算沿交线各点的打磨法向量，'
    '为机器人末端执行器的姿态规划和力控策略提供几何基础。'
)

doc.add_heading('今日完成', level=1)

doc.add_heading('1. 文献调研任务创建', level=2)
doc.add_paragraph('在飞书「科研任务」清单中创建"调研工业机器人末端力控相关文献"任务（截止 7月5日）。')

doc.add_heading('2. 正交圆柱交线几何建模', level=2)
doc.add_paragraph('完成两个正交圆柱相交曲线的通用数学推导，支持任意轴向组合（X/Y/Z），自动识别公共径向坐标。四段半曲线独立求解，自然汇合形成闭合空间曲线。')
p = doc.add_paragraph()
run = p.add_run('⚠ 当前数字为测试数据（单位mm），后续需金属零件详细参数替换。理想情况：如有实际测量的曲线数据，可用于验证模型精度。')
run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
run.bold = True

doc.add_heading('3. 打磨法向量计算', level=2)
doc.add_paragraph('基于两个圆柱面径向法向量的角平分线方向。物理场景：圆柱内部为空气（孔洞），外部为金属，法向量指向金属区域。方案已定，无需补充。')

doc.add_heading('4. 可视化代码实现', level=2)
doc.add_paragraph('核心代码：test.py，依赖 numpy 和 matplotlib。可配置参数（半径、位置、轴线方向）集中于文件顶部，输出 3D 可视化图像（含圆柱面、交线、法向量箭头、轴线）。')

doc.add_heading('技术方案', level=1)

doc.add_heading('误差计算', level=2)
doc.add_paragraph('需要先明确两个圆柱轴线的允许误差范围（位置偏差、角度偏差），然后：')
for item in ['在误差范围内取一组误差值', '计算包含误差的"真实"交线曲线', '与标准（无误差）曲线在同一图中对比', '同时计算包含误差后的法向量']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('误差修正算法', level=2)
doc.add_paragraph('基础算法：阻抗控制。核心思路：根据当前实际力测量值与预期值的差异，调整位置环的位置控制。控制器：先用 PID 控制，使末端贴合真实曲线。')

doc.add_heading('实验目标', level=2)
doc.add_paragraph('在软件仿真中验证算法有效性：在误差范围内生成随机误差曲线，运行修正算法。若能对误差范围内的任意曲线都能较好贴合，则算法有效。')

doc.add_heading('需要注意的问题', level=1)

table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['参数', '值', '分析']):
    table.rows[0].cells[i].text = h
    for r in table.rows[0].cells[i].paragraphs[0].runs:
        r.bold = True
for ri, row_data in enumerate([['平均测量力', '~8N', ''], ['传感器分辨率', '1N', '相对误差约 12.5%'], ['潜在风险', 'PID震荡', '力反馈步进变化（1N步长）可能导致位置环过度响应']]):
    for ci, val in enumerate(row_data):
        table.rows[ri+1].cells[ci].text = val

doc.add_paragraph('')
doc.add_heading('建议', level=2)
for s in ['力信号预处理 — 加入低通滤波或滑动窗口平均，平滑阶梯状力信号',
          'PID参数整定 — 建议先仿真跑一遍，观察是否出现极限环振荡',
          '考虑变增益 — 误差大时用大增益快速接近，误差小时用小增益避免震荡']:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('项目文件结构', level=1)
doc.add_paragraph('projects/formal/机器人末端力控/')
for item in ['project.md — 项目概述', '正交圆柱相交曲线分析.md — 完整数学建模文档', 'code/test.py — 可视化代码']:
    doc.add_paragraph(item, style='List Bullet')

doc.save('C:/Users/KCserver/projects/formal/机器人末端力控/progress.docx')
print('progress.docx created')
