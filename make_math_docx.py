from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

# Also set East Asian font
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = doc.element.makeelement(qn('w:rFonts'), {})
    rpr.append(rFonts)
rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_heading('正交圆柱相交曲线分析 & 打磨法向量计算', level=0)

doc.add_heading('1. 概述', level=1)
doc.add_paragraph(
    '本项目用 Python + matplotlib 对两个正交圆柱的相交曲线进行三维可视化分析，'
    '并计算曲线上每点的打磨法向量（指向金属区域）。适用于机械加工中圆柱孔相交边缘的打磨路径规划。'
)
doc.add_paragraph('核心文件：test.py')
doc.add_paragraph('输出图像：intersection_curve.png')
doc.add_paragraph('依赖：numpy、matplotlib')

doc.add_heading('2. 数学模型', level=1)

doc.add_heading('2.1 圆柱方程', level=2)
doc.add_paragraph(
    '设有两个圆柱，轴线分别平行于坐标轴，轴线方向不同（正交）。'
    '圆柱面的隐式方程（以轴线沿 X 为例）：(y - Cy)^2 + (z - Cz)^2 = R^2'
)

t = doc.add_table(rows=3, cols=3)
t.style = 'Light Grid Accent 1'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['参数', '圆柱 1', '圆柱 2']):
    t.rows[0].cells[i].text = h
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
t.rows[1].cells[0].text = '轴线方向'
t.rows[1].cells[1].text = 'AXIS1'
t.rows[1].cells[2].text = 'AXIS2'
t.rows[2].cells[0].text = '半径'
t.rows[2].cells[1].text = 'R1'
t.rows[2].cells[2].text = 'R2'

doc.add_heading('2.2 交线求解（通用正交情形）', level=2)
doc.add_paragraph(
    '对于轴线方向不同的两个圆柱，存在一个公共径向坐标（即同时作为两个圆柱径向的坐标轴）。'
    '通过集合运算自动确定各坐标的角色，从两个圆柱方程分别解出另一径向坐标。'
    '正负号组合产生 4 条半曲线，在端点处自然汇合，形成 1~2 条闭合空间曲线。'
)

doc.add_heading('2.3 打磨法向量', level=2)
doc.add_paragraph(
    '物理场景：圆柱内部 = 空气（被打穿的孔），圆柱外部 = 金属。'
    '相交曲线是金属-空气边界上的一条脊线，需沿此曲线打磨倒角。'
)
p = doc.add_paragraph()
p.add_run('计算公式：').bold = True
doc.add_paragraph('曲线上任一点 P 同时在两个圆柱面上，各自有向外的单位径向法向量 n1、n2。')
doc.add_paragraph('打磨法向量取二者的角平分线，指向金属（外侧）：n_grind = (n1 + n2) / |n1 + n2|')

doc.add_heading('3. 可配置参数', level=1)
doc.add_paragraph('所有参数集中在 test.py 顶部，修改后直接运行即可。')
doc.add_paragraph('圆柱几何参数：AXIS1/AXIS2（轴线方向）、CX/CY/CZ（轴线位置）、R1/R2（半径）、LEN1/LEN2（显示长度）')
doc.add_paragraph('显示参数：CYL_ALPHA（透明度）、AXIS_RANGE（坐标轴范围）、NORMAL_SAMPLE（法向量采样间隔）、NORMAL_LENGTH（箭头长度）')

doc.add_heading('4. 代码架构', level=1)
code_text = (
    "test.py\n"
    "├── 可配置参数区\n"
    "├── 工具函数\n"
    "│   ├── make_cylinder()  —— 生成任意轴向的圆柱面 meshgrid\n"
    "│   └── axis_line()      —— 生成轴线端点坐标\n"
    "├── 交线求解模块          —— 自动识别公共径向坐标，解出四段半曲线\n"
    "├── 法向量计算模块        —— 逐点计算两个圆柱面法向量 → 角平分线\n"
    "└── 绘图输出              —— 3D 可视化 + 图片保存"
)
doc.add_paragraph(code_text)

doc.add_heading('关键设计决策', level=2)
items = [
    '通用轴向支持：基于坐标索引的集合运算，自动适配 X/Y/Z 任意轴向组合',
    '四段独立绘制：避免拼接顺序错误导致的虚假连线',
    '先保存再显示：savefig 必须在 show 之前调用',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5. 测试过的参数组合', level=1)
t2 = doc.add_table(rows=4, cols=6)
t2.style = 'Light Grid Accent 1'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['场景', 'AXIS1', 'AXIS2', 'R1', 'R2', '交线特征']
for i, h in enumerate(headers):
    t2.rows[0].cells[i].text = h
    t2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
data = [
    ['初始测试', 'X', 'Z', '10', '20', '1 条闭合曲线'],
    ['验证轴向切换', 'Y', 'Z', '10', '20', '1 条闭合曲线'],
    ['中心偏移', 'Y', 'Z', '10', '20', '验证法向量方向'],
]
for ri, row in enumerate(data):
    for ci, val in enumerate(row):
        t2.rows[ri+1].cells[ci].text = val

doc.add_paragraph('')

doc.add_heading('6. 踩坑记录', level=1)
t3 = doc.add_table(rows=5, cols=3)
t3.style = 'Light Grid Accent 1'
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['问题', '原因', '解决']):
    t3.rows[0].cells[i].text = h
    t3.rows[0].cells[i].paragraphs[0].runs[0].bold = True
pitfalls = [
    ['相交曲线出现多余连线', '四段拼接顺序不对', '改为四段各自独立 plot()'],
    ['保存图像为空白', 'show() 后图像清空', 'savefig 移到 show 之前'],
    ['三轴比例不一致', 'box_aspect 只改包围盒', '手动设三轴相同跨度'],
    ['轴向硬编码', '假设固定轴向', '基于坐标索引的通用函数'],
]
for ri, row in enumerate(pitfalls):
    for ci, val in enumerate(row):
        t3.rows[ri+1].cells[ci].text = val

doc.add_heading('7. 运行方式', level=1)
doc.add_paragraph('cd e:\\Code\\Claude\\ForceFeedback')
doc.add_paragraph('python test.py')
doc.add_paragraph('运行后弹出交互式 3D 窗口（可拖拽旋转/缩放），关闭窗口后图像自动保存到 intersection_curve.png。')

doc.save('C:/Users/KCserver/projects/formal/机器人末端力控/正交圆柱相交曲线分析.docx')
print('正交圆柱相交曲线分析.docx created')
